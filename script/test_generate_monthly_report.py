"""Tests for generate_monthly_report.py using unittest."""

from __future__ import annotations

import io
import json
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

# Add script directory to path so we can import the module under test
sys.path.insert(0, str(Path(__file__).parent))

import generate_monthly_report as gmr
from docx import Document


TEMPLATE_PATH = Path(__file__).parent / "templates" / "monthly_report_template.docx"


def _make_temp_output() -> Path:
    """Return a temporary file path in an existing directory for output DOCX."""
    tmp_dir = Path(tempfile.mkdtemp())
    return tmp_dir / "output.docx"


def _run_main(data: dict) -> Path:
    """Run main() with the given data dict, return the output path."""
    output_path = _make_temp_output()
    data = {**data, "output_path": str(output_path)}
    stdin_text = json.dumps(data)
    with patch("sys.stdin", io.StringIO(stdin_text)):
        gmr.main()
    return output_path


class TestFormatDateKorean(unittest.TestCase):
    def test_basic_conversion(self) -> None:
        self.assertEqual(gmr.format_date_korean("2026-02-11"), "2026년 2월 11일")

    def test_single_digit_month_and_day(self) -> None:
        self.assertEqual(gmr.format_date_korean("2026-01-05"), "2026년 1월 5일")

    def test_december(self) -> None:
        self.assertEqual(gmr.format_date_korean("2025-12-31"), "2025년 12월 31일")


class TestFormatSubmissionDateSpaced(unittest.TestCase):
    def test_basic_conversion(self) -> None:
        self.assertEqual(
            gmr.format_submission_date_spaced("2026-03-16"),
            "2026 년   3 월 16일",
        )

    def test_single_digit_month(self) -> None:
        result = gmr.format_submission_date_spaced("2026-01-10")
        self.assertIn("2026", result)
        self.assertIn("1", result)
        self.assertIn("10일", result)


class TestValidateInput(unittest.TestCase):
    def _base_data(self) -> dict:
        return {
            "activities": [
                {"date": "2026-02-11", "description": "도서 리뷰", "count": 2}
            ],
            "total": 2,
            "submission_date": "2026-03-18",
            "output_path": str(_make_temp_output()),
        }

    def test_valid_input_passes(self) -> None:
        data = self._base_data()
        # Should not raise
        gmr.validate_input(data)

    def test_missing_activities_raises(self) -> None:
        data = self._base_data()
        del data["activities"]
        with self.assertRaises(ValueError) as ctx:
            gmr.validate_input(data)
        self.assertIn("activities", str(ctx.exception))

    def test_missing_total_raises(self) -> None:
        data = self._base_data()
        del data["total"]
        with self.assertRaises(ValueError) as ctx:
            gmr.validate_input(data)
        self.assertIn("total", str(ctx.exception))

    def test_missing_submission_date_raises(self) -> None:
        data = self._base_data()
        del data["submission_date"]
        with self.assertRaises(ValueError) as ctx:
            gmr.validate_input(data)
        self.assertIn("submission_date", str(ctx.exception))

    def test_missing_output_path_raises(self) -> None:
        data = self._base_data()
        del data["output_path"]
        with self.assertRaises(ValueError) as ctx:
            gmr.validate_input(data)
        self.assertIn("output_path", str(ctx.exception))

    def test_more_than_8_activities_raises(self) -> None:
        data = self._base_data()
        data["activities"] = [
            {"date": f"2026-02-{i+1:02d}", "description": f"activity {i}", "count": 1}
            for i in range(9)
        ]
        with self.assertRaises(ValueError) as ctx:
            gmr.validate_input(data)
        self.assertIn("Too many activities", str(ctx.exception))

    def test_exactly_8_activities_passes(self) -> None:
        output_path = _make_temp_output()
        data = {
            "activities": [
                {"date": f"2026-02-{i+1:02d}", "description": f"act {i}", "count": 1}
                for i in range(8)
            ],
            "total": 8,
            "submission_date": "2026-03-18",
            "output_path": str(output_path),
        }
        # Should not raise
        gmr.validate_input(data)

    def test_activity_missing_date_raises(self) -> None:
        data = self._base_data()
        data["activities"] = [{"description": "no date", "count": 1}]
        with self.assertRaises(ValueError) as ctx:
            gmr.validate_input(data)
        self.assertIn("date", str(ctx.exception))

    def test_activity_missing_description_raises(self) -> None:
        data = self._base_data()
        data["activities"] = [{"date": "2026-02-01", "count": 1}]
        with self.assertRaises(ValueError) as ctx:
            gmr.validate_input(data)
        self.assertIn("description", str(ctx.exception))

    def test_activity_missing_count_raises(self) -> None:
        data = self._base_data()
        data["activities"] = [{"date": "2026-02-01", "description": "no count"}]
        with self.assertRaises(ValueError) as ctx:
            gmr.validate_input(data)
        self.assertIn("count", str(ctx.exception))

    def test_nonexistent_output_directory_raises(self) -> None:
        data = self._base_data()
        data["output_path"] = "/nonexistent/path/output.docx"
        with self.assertRaises(ValueError) as ctx:
            gmr.validate_input(data)
        self.assertIn("Output directory", str(ctx.exception))


class TestNormalCase(unittest.TestCase):
    SAMPLE_ACTIVITIES = [
        {"date": "2026-02-11", "description": "도서 리뷰 (동탄)", "count": 2},
        {"date": "2026-02-18", "description": "독서 토론 (수원)", "count": 4},
    ]

    def test_output_file_created(self) -> None:
        output_path = _run_main(
            {
                "activities": self.SAMPLE_ACTIVITIES,
                "total": 6,
                "submission_date": "2026-03-18",
            }
        )
        self.assertTrue(output_path.exists(), "Output DOCX was not created")

    def test_output_is_valid_docx(self) -> None:
        output_path = _run_main(
            {
                "activities": self.SAMPLE_ACTIVITIES,
                "total": 6,
                "submission_date": "2026-03-18",
            }
        )
        # Document() raises if the file is not a valid DOCX
        doc = Document(str(output_path))
        self.assertIsNotNone(doc)

    def test_activity_rows_written(self) -> None:
        output_path = _run_main(
            {
                "activities": self.SAMPLE_ACTIVITIES,
                "total": 6,
                "submission_date": "2026-03-18",
            }
        )
        doc = Document(str(output_path))
        table = doc.tables[0]

        # Row 6 (index 6): first activity
        row6 = table.rows[6]
        self.assertIn("2026년", row6.cells[gmr.DATE_CELL_INDEX].text)
        self.assertIn("2월", row6.cells[gmr.DATE_CELL_INDEX].text)
        self.assertIn("11일", row6.cells[gmr.DATE_CELL_INDEX].text)
        self.assertIn("도서 리뷰 (동탄)", row6.cells[gmr.DESC_CELL_INDEX].text)
        self.assertIn("2명", row6.cells[gmr.COUNT_CELL_INDEX].text)

        # Row 7 (index 7): second activity
        row7 = table.rows[7]
        self.assertIn("18일", row7.cells[gmr.DATE_CELL_INDEX].text)
        self.assertIn("독서 토론 (수원)", row7.cells[gmr.DESC_CELL_INDEX].text)
        self.assertIn("4명", row7.cells[gmr.COUNT_CELL_INDEX].text)

    def test_unused_rows_cleared(self) -> None:
        output_path = _run_main(
            {
                "activities": self.SAMPLE_ACTIVITIES,  # only 2 activities
                "total": 6,
                "submission_date": "2026-03-18",
            }
        )
        doc = Document(str(output_path))
        table = doc.tables[0]

        # Rows 8-13 (indices 8..13) should be cleared
        for row_idx in range(8, 14):
            row = table.rows[row_idx]
            for cell_idx in (gmr.DATE_CELL_INDEX, gmr.DESC_CELL_INDEX, gmr.COUNT_CELL_INDEX):
                self.assertEqual(
                    row.cells[cell_idx].text.strip(),
                    "",
                    f"Row {row_idx} cell {cell_idx} should be empty",
                )

    def test_total_written_with_suffix(self) -> None:
        output_path = _run_main(
            {
                "activities": self.SAMPLE_ACTIVITIES,
                "total": 6,
                "submission_date": "2026-03-18",
            }
        )
        doc = Document(str(output_path))
        table = doc.tables[0]
        total_cell_text = table.rows[14].cells[gmr.COUNT_CELL_INDEX].text
        self.assertIn("6명", total_cell_text)

    def test_submission_date_written(self) -> None:
        output_path = _run_main(
            {
                "activities": self.SAMPLE_ACTIVITIES,
                "total": 6,
                "submission_date": "2026-03-18",
            }
        )
        doc = Document(str(output_path))
        table = doc.tables[0]
        row16_cell = table.rows[16].cells[0]
        para_text = row16_cell.paragraphs[gmr.SUBMISSION_DATE_PARA_INDEX].text
        self.assertIn("2026", para_text)
        self.assertIn("3", para_text)
        self.assertIn("18일", para_text)


class TestEmptyActivities(unittest.TestCase):
    def test_empty_activities_all_rows_cleared(self) -> None:
        output_path = _run_main(
            {
                "activities": [],
                "total": 0,
                "submission_date": "2026-03-18",
            }
        )
        doc = Document(str(output_path))
        table = doc.tables[0]

        for row_idx in range(6, 14):
            row = table.rows[row_idx]
            for cell_idx in (gmr.DATE_CELL_INDEX, gmr.DESC_CELL_INDEX, gmr.COUNT_CELL_INDEX):
                self.assertEqual(
                    row.cells[cell_idx].text.strip(),
                    "",
                    f"Row {row_idx} cell {cell_idx} should be empty for zero activities",
                )

    def test_empty_activities_total_zero(self) -> None:
        output_path = _run_main(
            {
                "activities": [],
                "total": 0,
                "submission_date": "2026-03-18",
            }
        )
        doc = Document(str(output_path))
        table = doc.tables[0]
        total_text = table.rows[14].cells[gmr.COUNT_CELL_INDEX].text
        self.assertIn("0명", total_text)

    def test_output_is_valid_docx(self) -> None:
        output_path = _run_main(
            {
                "activities": [],
                "total": 0,
                "submission_date": "2026-03-18",
            }
        )
        doc = Document(str(output_path))
        self.assertIsNotNone(doc)


class TestMaxActivities(unittest.TestCase):
    def _eight_activities(self) -> list[dict]:
        return [
            {"date": f"2026-02-{i+1:02d}", "description": f"활동 {i+1}", "count": i + 1}
            for i in range(8)
        ]

    def test_exactly_8_activities_all_written(self) -> None:
        acts = self._eight_activities()
        output_path = _run_main(
            {
                "activities": acts,
                "total": sum(a["count"] for a in acts),
                "submission_date": "2026-03-18",
            }
        )
        doc = Document(str(output_path))
        table = doc.tables[0]

        for i in range(8):
            row = table.rows[6 + i]
            self.assertIn(
                acts[i]["description"],
                row.cells[gmr.DESC_CELL_INDEX].text,
                f"Activity {i} description not found in row {6+i}",
            )

    def test_9_activities_raises_value_error(self) -> None:
        acts = [
            {"date": f"2026-02-{i+1:02d}", "description": f"활동 {i+1}", "count": 1}
            for i in range(9)
        ]
        output_path = _make_temp_output()
        stdin_text = json.dumps(
            {
                "activities": acts,
                "total": 9,
                "submission_date": "2026-03-18",
                "output_path": str(output_path),
            }
        )
        with patch("sys.stdin", io.StringIO(stdin_text)):
            with self.assertRaises(ValueError) as ctx:
                gmr.main()
        self.assertIn("Too many activities", str(ctx.exception))


class TestDateSorting(unittest.TestCase):
    def test_unsorted_input_sorted_in_output(self) -> None:
        # Provide activities in reverse order — they should appear sorted ASC in the DOCX
        activities = [
            {"date": "2026-02-25", "description": "세 번째", "count": 3},
            {"date": "2026-02-05", "description": "첫 번째", "count": 1},
            {"date": "2026-02-15", "description": "두 번째", "count": 2},
        ]
        output_path = _run_main(
            {
                "activities": activities,
                "total": 6,
                "submission_date": "2026-03-18",
            }
        )
        doc = Document(str(output_path))
        table = doc.tables[0]

        row6_desc = table.rows[6].cells[gmr.DESC_CELL_INDEX].text
        row7_desc = table.rows[7].cells[gmr.DESC_CELL_INDEX].text
        row8_desc = table.rows[8].cells[gmr.DESC_CELL_INDEX].text

        self.assertIn("첫 번째", row6_desc, "Earliest date should be in row 6")
        self.assertIn("두 번째", row7_desc, "Middle date should be in row 7")
        self.assertIn("세 번째", row8_desc, "Latest date should be in row 8")

    def test_already_sorted_input_unchanged_order(self) -> None:
        activities = [
            {"date": "2026-02-01", "description": "처음", "count": 1},
            {"date": "2026-02-28", "description": "마지막", "count": 2},
        ]
        output_path = _run_main(
            {
                "activities": activities,
                "total": 3,
                "submission_date": "2026-03-18",
            }
        )
        doc = Document(str(output_path))
        table = doc.tables[0]

        self.assertIn("처음", table.rows[6].cells[gmr.DESC_CELL_INDEX].text)
        self.assertIn("마지막", table.rows[7].cells[gmr.DESC_CELL_INDEX].text)


def _create_tiny_png(path: Path) -> None:
    """Create a minimal valid 1x1 white PNG file."""
    import struct
    import zlib

    def _chunk(chunk_type: bytes, data: bytes) -> bytes:
        c = chunk_type + data
        crc = struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + c + crc

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = _chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
    raw_row = b"\x00\xff\xff\xff"  # filter byte + RGB
    idat = _chunk(b"IDAT", zlib.compress(raw_row))
    iend = _chunk(b"IEND", b"")
    path.write_bytes(sig + ihdr + idat + iend)


class TestPhotoPagesBackwardCompat(unittest.TestCase):
    """Photos field is optional — existing inputs without it must work identically."""

    SAMPLE = {
        "activities": [
            {"date": "2026-02-11", "description": "도서 리뷰", "count": 2},
        ],
        "total": 2,
        "submission_date": "2026-03-18",
    }

    def test_no_photos_field(self) -> None:
        output_path = _run_main(self.SAMPLE)
        doc = Document(str(output_path))
        self.assertIsNotNone(doc)
        # Only the original table content, no extra paragraphs with page breaks
        self.assertEqual(len(doc.tables), 1)

    def test_empty_photos_list(self) -> None:
        data = {**self.SAMPLE, "photos": []}
        output_path = _run_main(data)
        doc = Document(str(output_path))
        self.assertEqual(len(doc.tables), 1)


class TestPhotoPages(unittest.TestCase):
    """Test attendance photo page generation."""

    SAMPLE_BASE = {
        "activities": [
            {"date": "2026-02-11", "description": "도서 리뷰", "count": 2},
        ],
        "total": 2,
        "submission_date": "2026-03-18",
    }

    def test_local_file_photo_embedded(self) -> None:
        tmp_dir = Path(tempfile.mkdtemp())
        img_path = tmp_dir / "test_photo.png"
        _create_tiny_png(img_path)

        data = {
            **self.SAMPLE_BASE,
            "photos": [
                {
                    "date": "2026-02-11",
                    "meeting_title": "도서 리뷰 (동탄)",
                    "file_path": str(img_path),
                    "caption": "단체 사진",
                },
            ],
        }
        output_path = _run_main(data)
        doc = Document(str(output_path))

        # Should have at least one inline shape (the embedded image)
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        image_parts = [
            rel.target_ref
            for rel in doc.part.rels.values()
            if "image" in rel.reltype
        ]
        self.assertGreater(len(image_parts), 0, "No image found in document")

    def test_heading_contains_korean_date_and_title(self) -> None:
        tmp_dir = Path(tempfile.mkdtemp())
        img_path = tmp_dir / "test_photo.png"
        _create_tiny_png(img_path)

        data = {
            **self.SAMPLE_BASE,
            "photos": [
                {
                    "date": "2026-02-11",
                    "meeting_title": "도서 리뷰 (동탄)",
                    "file_path": str(img_path),
                },
            ],
        }
        output_path = _run_main(data)
        doc = Document(str(output_path))

        # Find the bold heading paragraph
        found = False
        for para in doc.paragraphs:
            if para.runs and para.runs[0].bold:
                text = para.text
                if "2026년 2월 11일" in text and "도서 리뷰 (동탄)" in text:
                    found = True
                    break
        self.assertTrue(found, "Bold heading with Korean date and meeting title not found")

    def test_caption_appears_in_document(self) -> None:
        tmp_dir = Path(tempfile.mkdtemp())
        img_path = tmp_dir / "test_photo.png"
        _create_tiny_png(img_path)

        data = {
            **self.SAMPLE_BASE,
            "photos": [
                {
                    "date": "2026-02-11",
                    "meeting_title": "도서 리뷰",
                    "file_path": str(img_path),
                    "caption": "참석자 단체 사진",
                },
            ],
        }
        output_path = _run_main(data)
        doc = Document(str(output_path))

        all_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("참석자 단체 사진", all_text)

    def test_missing_file_warns_no_crash(self) -> None:
        data = {
            **self.SAMPLE_BASE,
            "photos": [
                {
                    "date": "2026-02-11",
                    "meeting_title": "도서 리뷰",
                    "file_path": "/nonexistent/photo.jpg",
                },
            ],
        }
        import io as _io

        captured = _io.StringIO()
        with patch("sys.stderr", captured):
            output_path = _run_main(data)

        self.assertTrue(output_path.exists())
        self.assertIn("Warning", captured.getvalue())
        self.assertIn("/nonexistent/photo.jpg", captured.getvalue())

    def test_photos_not_a_list_raises(self) -> None:
        data = {**self.SAMPLE_BASE, "photos": "not_a_list"}
        output_path = _make_temp_output()
        data["output_path"] = str(output_path)
        with self.assertRaises(ValueError) as ctx:
            gmr.validate_input(data)
        self.assertIn("photos must be a list", str(ctx.exception))


if __name__ == "__main__":
    unittest.main()
