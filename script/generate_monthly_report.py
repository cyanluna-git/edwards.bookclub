"""Generate monthly report DOCX from JSON input via stdin.

Input JSON schema:
{
  "activities": [
    {"date": "2026-02-11", "description": "도서 리뷰 (동탄)", "count": 2},
    ...
  ],
  "total": 23,
  "submission_date": "2026-03-16",
  "output_path": "/path/to/output.docx",
  "photos": [
    {
      "date": "2026-02-11",
      "meeting_title": "도서 리뷰 (동탄)",
      "file_path": "/path/to/local/photo.jpg",
      "source_url": "https://example.com/photo.jpg",
      "caption": "단체 사진"
    },
    ...
  ]
}

- date format: YYYY-MM-DD (ISO 8601)
- description: activity description with optional location in parentheses
- count: integer number of participants
- total: integer total participants across all activities
- submission_date: YYYY-MM-DD
- output_path: absolute or relative path for output file
- photos (optional): list of attendance photo objects
  - date: YYYY-MM-DD
  - meeting_title: title for the photo heading
  - file_path (optional): local file path to the image
  - source_url (optional): HTTP URL to fetch the image from
  - caption (optional): caption text below the image
"""

from __future__ import annotations

import io
import json
import sys
from pathlib import Path
from typing import Any

import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.table import _Cell
from docx.text.paragraph import Paragraph

TEMPLATE_PATH = Path(__file__).parent / "templates" / "monthly_report_template.docx"

ACTIVITY_ROW_START = 6
ACTIVITY_ROW_END = 13  # inclusive
MAX_ACTIVITIES = ACTIVITY_ROW_END - ACTIVITY_ROW_START + 1  # 8

DATE_CELL_INDEX = 2
DESC_CELL_INDEX = 6
COUNT_CELL_INDEX = 14
SUBMISSION_DATE_PARA_INDEX = 2  # 0-based within row 16 cell


def set_cell_text(cell: _Cell, text: str) -> None:
    """Clear all runs in the first paragraph and set new text, preserving paragraph properties."""
    paragraph: Paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def format_date_korean(date_str: str) -> str:
    """Convert ISO date string (YYYY-MM-DD) to Korean format without zero-padding.

    Example: "2026-02-11" -> "2026년 2월 11일"
    """
    parts = date_str.split("-")
    year = int(parts[0])
    month = int(parts[1])
    day = int(parts[2])
    return f"{year}년 {month}월 {day}일"


def format_submission_date_spaced(date_str: str) -> str:
    """Convert ISO date string to spaced Korean format matching template style.

    Example: "2026-03-16" -> "2026 년   3 월 16일"
    """
    parts = date_str.split("-")
    year = int(parts[0])
    month = int(parts[1])
    day = int(parts[2])
    return f"{year} 년   {month} 월 {day}일"


def write_activities(table: Any, activities: list[dict[str, Any]]) -> None:
    """Write activity data to rows 6-13. Unused rows are cleared."""
    for i in range(MAX_ACTIVITIES):
        row_index = ACTIVITY_ROW_START + i
        row = table.rows[row_index]

        if i < len(activities):
            activity = activities[i]
            set_cell_text(row.cells[DATE_CELL_INDEX], format_date_korean(activity["date"]))
            set_cell_text(row.cells[DESC_CELL_INDEX], activity["description"])
            set_cell_text(row.cells[COUNT_CELL_INDEX], f"{activity['count']}명")
        else:
            set_cell_text(row.cells[DATE_CELL_INDEX], "")
            set_cell_text(row.cells[DESC_CELL_INDEX], "")
            set_cell_text(row.cells[COUNT_CELL_INDEX], "")


def write_total(table: Any, total: int) -> None:
    """Write total participant count to row 14, cells[14]."""
    set_cell_text(table.rows[14].cells[COUNT_CELL_INDEX], f"{total}명")


def write_submission_date(table: Any, date_str: str) -> None:
    """Write submission date to row 16, paragraph index 2 (0-based)."""
    cell = table.rows[16].cells[0]
    paragraph = cell.paragraphs[SUBMISSION_DATE_PARA_INDEX]
    formatted = format_submission_date_spaced(date_str)

    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = formatted
    else:
        paragraph.add_run(formatted)


def fetch_image(photo: dict[str, Any]) -> io.BytesIO | Path | None:
    """Resolve a photo to a local Path or BytesIO from HTTP. Returns None on failure."""
    file_path = photo.get("file_path")
    if file_path:
        p = Path(file_path)
        if p.is_file():
            return p
        print(f"Warning: photo file not found: {file_path}", file=sys.stderr)
        return None

    source_url = photo.get("source_url")
    if source_url:
        try:
            resp = requests.get(source_url, timeout=30)
            resp.raise_for_status()
            buf = io.BytesIO(resp.content)
            return buf
        except Exception as exc:
            print(f"Warning: failed to fetch photo from {source_url}: {exc}", file=sys.stderr)
            return None

    print("Warning: photo has neither file_path nor source_url, skipping", file=sys.stderr)
    return None


def write_photo_pages(doc: Document, photos: list[dict[str, Any]]) -> None:
    """Append attendance photo pages to the document."""
    for photo in photos:
        image = fetch_image(photo)
        if image is None:
            continue

        doc.add_page_break()

        date_str = photo.get("meeting_date", "") or photo.get("date", "")
        meeting_title = photo.get("meeting_title", "")
        location = photo.get("location", "")
        attendees = photo.get("attendees", "")

        heading_text = f"{format_date_korean(date_str)} {meeting_title}".strip() if date_str else meeting_title
        heading_para = doc.add_paragraph()
        run = heading_para.add_run(heading_text)
        run.bold = True

        if location:
            loc_para = doc.add_paragraph()
            loc_run = loc_para.add_run(f"장소: {location}")
            loc_run.font.size = Pt(10)

        if attendees:
            att_para = doc.add_paragraph()
            att_run = att_para.add_run(f"참석자: {attendees}")
            att_run.font.size = Pt(10)

        image_arg = str(image) if isinstance(image, Path) else image
        doc.add_picture(image_arg, width=Inches(6.5))

        caption = photo.get("caption")
        if caption:
            doc.add_paragraph(caption)


def validate_input(data: dict[str, Any]) -> None:
    """Validate input JSON data, raising ValueError on invalid input."""
    if "activities" not in data:
        raise ValueError("Missing required field: activities")
    if "total" not in data:
        raise ValueError("Missing required field: total")
    if "submission_date" not in data:
        raise ValueError("Missing required field: submission_date")
    if "output_path" not in data:
        raise ValueError("Missing required field: output_path")

    activities = data["activities"]
    if not isinstance(activities, list):
        raise ValueError("activities must be a list")
    if len(activities) > MAX_ACTIVITIES:
        raise ValueError(
            f"Too many activities: {len(activities)} (max {MAX_ACTIVITIES})"
        )

    for i, act in enumerate(activities):
        for field in ("date", "description", "count"):
            if field not in act:
                raise ValueError(f"Activity {i}: missing required field '{field}'")

    photos = data.get("photos")
    if photos is not None and not isinstance(photos, list):
        raise ValueError("photos must be a list")

    output_dir = Path(data["output_path"]).parent
    if not output_dir.exists():
        raise ValueError(f"Output directory does not exist: {output_dir}")


def main() -> None:
    """Read JSON from stdin, fill template, write DOCX to output_path."""
    raw = sys.stdin.read()
    data: dict[str, Any] = json.loads(raw)

    validate_input(data)

    activities = sorted(data["activities"], key=lambda a: a["date"])

    doc = Document(str(TEMPLATE_PATH))
    table = doc.tables[0]

    write_activities(table, activities)
    write_total(table, data["total"])
    write_submission_date(table, data["submission_date"])

    photos = data.get("photos", [])
    if photos:
        write_photo_pages(doc, photos)

    output_path = Path(data["output_path"])
    doc.save(str(output_path))
    print(f"Report saved to {output_path}")


if __name__ == "__main__":
    main()
